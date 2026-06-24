"""
Validador de documentos Word para planificaciones.
Verifica que el documento contenga las secciones indicadas en el TipoPlanificacion.
Las secciones son texto libre configurado por la moderadora (ej: "Propósito", "Bibliografía").
"""
from docx import Document


def validar_documento_word(archivo, campos_requeridos=None):
    """
    Valida que el documento Word contenga las secciones indicadas.

    Args:
        archivo: objeto file-like (FileField abierto o BytesIO)
        campos_requeridos: lista de strings a buscar en el documento (texto libre).
                           Si es None o vacío, no se valida nada → siempre válido.

    Returns:
        (es_valido: bool, campos_faltantes: list[str])
    """
    if not campos_requeridos:
        return True, []

    try:
        doc = Document(archivo)
    except Exception as e:
        return False, [f'Error al leer el documento: {e}']

    # Extraer todo el texto del documento en minúsculas
    textos = []
    for para in doc.paragraphs:
        textos.append(para.text.lower())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                textos.append(cell.text.lower())
    texto = ' '.join(textos)

    # Buscar cada sección requerida (búsqueda case-insensitive)
    campos_faltantes = []
    for campo in campos_requeridos:
        if campo.strip().lower() not in texto:
            campos_faltantes.append(campo)

    return len(campos_faltantes) == 0, campos_faltantes


def nombre_campo(campo):
    """Retorna el nombre de un campo (compatibilidad — con texto libre, ya es el nombre)."""
    return campo
