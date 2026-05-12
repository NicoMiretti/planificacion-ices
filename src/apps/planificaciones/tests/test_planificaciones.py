"""
Tests para el módulo de planificaciones.
Verifica: validador Word, modelos, FSM, entregas tardías, vistas.
Ejecutar: docker-compose exec web pytest apps/planificaciones/tests/ -v
"""
import pytest
from io import BytesIO
from datetime import date, timedelta

from django.test import Client
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document

from apps.catalogos.models import Institucion, Carrera, Materia, Profesor
from apps.instancias.models import InstanciaPresentacion
from apps.planificaciones.models import Planificacion, Version
from apps.planificaciones.validators import validar_documento_word, nombre_campo
from apps.usuarios.models import Usuario


# ============================================================
# Fixtures comunes
# ============================================================

@pytest.fixture
def institucion():
    return Institucion.objects.create(nombre='ICES Test', codigo='ICEST')


@pytest.fixture
def carrera(institucion):
    return Carrera.objects.create(nombre='Ingeniería en Sistemas', institucion=institucion)


@pytest.fixture
def usuario_profesor():
    return Usuario.objects.create_user(
        email='prof.plan@ices.edu',
        password='test123',
        rol=Usuario.Rol.PROFESOR,
        nombre_completo='Prof Planificaciones'
    )


@pytest.fixture
def profesor(usuario_profesor, institucion):
    return Profesor.objects.create(
        usuario=usuario_profesor,
        institucion=institucion,
        legajo='P-001'
    )


@pytest.fixture
def materia(carrera, profesor):
    return Materia.objects.create(
        nombre='Programación I',
        carrera=carrera,
        anio_cursado=1,
        regimen=Materia.Regimen.ANUAL,
        profesor_titular=profesor
    )


@pytest.fixture
def instancia(carrera):
    inst = InstanciaPresentacion.objects.create(
        nombre='Planificaciones 2026',
        anio_academico=2026,
        periodo=InstanciaPresentacion.Periodo.ANUAL,
        fecha_apertura=date.today() - timedelta(days=5),
        fecha_limite=date.today() + timedelta(days=25),
    )
    inst.carreras.add(carrera)
    return inst


@pytest.fixture
def planificacion(materia, profesor, instancia):
    p, _ = Planificacion.objects.get_or_create(
        materia=materia,
        profesor=profesor,
        instancia=instancia
    )
    return p


def crear_doc_word(campos):
    """Crea un documento Word en memoria con los campos indicados."""
    doc = Document()
    for campo in campos:
        doc.add_heading(campo, level=1)
        doc.add_paragraph(f'Contenido de {campo}.')
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def doc_completo():
    """Documento Word con los 7 campos obligatorios."""
    return crear_doc_word([
        'Propósito',
        'Fundamentación',
        'Contenidos Mínimos',
        'Metodología',
        'Requisitos para regularizar',
        'Criterios de evaluación',
        'Bibliografía',
    ])


def doc_incompleto():
    """Documento Word con solo 2 campos."""
    return crear_doc_word(['Propósito', 'Bibliografía'])


# ============================================================
# TESTS DEL VALIDADOR
# ============================================================

class TestValidador:
    """Tests del validador de documentos Word (sin DB)."""

    def test_documento_completo_es_valido(self):
        """Documento con los 7 campos pasa la validación."""
        es_valido, faltantes = validar_documento_word(doc_completo())
        assert es_valido is True
        assert faltantes == []

    def test_documento_incompleto_no_es_valido(self):
        """Documento con campos faltantes falla la validación."""
        es_valido, faltantes = validar_documento_word(doc_incompleto())
        assert es_valido is False
        assert len(faltantes) == 5  # Faltan 5 de 7

    def test_documento_vacio_falla(self):
        """Documento completamente vacío falla con todos los campos faltantes."""
        doc = Document()
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        es_valido, faltantes = validar_documento_word(buf)
        assert es_valido is False
        assert len(faltantes) == 7

    def test_nombre_campo_devuelve_string(self):
        """nombre_campo retorna string legible."""
        assert nombre_campo('proposito') == 'Propósito'
        assert nombre_campo('bibliografia') == 'Bibliografía'

    def test_variantes_de_campo(self):
        """Acepta variantes de texto para el mismo campo."""
        doc = crear_doc_word([
            'objetivo general',   # variante de 'proposito'
            'fundamentacion',     # sin tilde
            'contenidos minimos', # sin tilde
            'estrategia',         # variante de 'metodologia'
            'regularizar',        # variante de 'requisitos'
            'evaluacion',         # sin tilde
            'bibliografia',       # sin tilde
        ])
        es_valido, faltantes = validar_documento_word(doc)
        assert es_valido is True
        assert faltantes == []


# ============================================================
# TESTS DEL MODELO PLANIFICACION
# ============================================================

@pytest.mark.django_db
class TestPlanificacionModelo:
    """Tests del modelo Planificacion."""

    def test_crear_planificacion(self, materia, profesor, instancia):
        """RF-07: Se puede crear una planificación para materia/profesor/instancia."""
        planif, _ = Planificacion.objects.get_or_create(
            materia=materia,
            profesor=profesor,
            instancia=instancia
        )
        assert planif.pk is not None
        assert str(planif) == f"{materia.nombre} - {profesor} ({instancia})"

    def test_unique_together(self, planificacion, materia, profesor, instancia):
        """No se puede crear dos planificaciones para la misma materia/profesor/instancia."""
        with pytest.raises(Exception):
            Planificacion.objects.create(
                materia=materia,
                profesor=profesor,
                instancia=instancia
            )

    def test_siguiente_numero_version_sin_versiones(self, planificacion):
        """Sin versiones, el siguiente número es 1."""
        assert planificacion.siguiente_numero_version() == 1

    def test_siguiente_numero_version_con_versiones(self, planificacion):
        """Con versiones existentes, incrementa correctamente."""
        archivo = SimpleUploadedFile('test.docx', b'contenido', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        Version.objects.create(planificacion=planificacion, numero=1, archivo=archivo)
        assert planificacion.siguiente_numero_version() == 2

    def test_no_tiene_oficial_inicialmente(self, planificacion):
        """Nueva planificación no tiene versión oficial."""
        assert planificacion.tiene_oficial is False

    def test_ultima_version_none_sin_versiones(self, planificacion):
        """Sin versiones, ultima_version devuelve None."""
        assert planificacion.ultima_version is None


# ============================================================
# TESTS DE LA FSM (VERSIONES)
# ============================================================

@pytest.mark.django_db
class TestVersionFSM:
    """Tests de las transiciones de estado de Version."""

    @pytest.fixture
    def version(self, planificacion):
        archivo = SimpleUploadedFile('test.docx', b'contenido', content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return Version.objects.create(planificacion=planificacion, numero=1, archivo=archivo)

    def test_estado_inicial_borrador(self, version):
        """Una versión nueva empieza en estado borrador."""
        assert version.estado == Version.Estado.BORRADOR

    def test_enviar_cambia_estado(self, version):
        """enviar() cambia estado a ENVIADA."""
        version.enviar()
        version.save()
        assert version.estado == Version.Estado.ENVIADA
        assert version.fecha_envio is not None

    def test_rechazar_automaticamente(self, version):
        """rechazar_automaticamente() cambia estado a RECHAZADA_AUTO."""
        campos = ['proposito', 'fundamentacion']
        version.rechazar_automaticamente(campos)
        version.save()
        assert version.estado == Version.Estado.RECHAZADA_AUTO
        assert version.campos_faltantes == campos

    def test_flujo_completo_aprobacion(self, version):
        """Flujo: borrador → enviada → en_revision → aprobada."""
        version.enviar()
        version.save()
        assert version.estado == Version.Estado.ENVIADA

        version.tomar_revision()
        version.save()
        assert version.estado == Version.Estado.EN_REVISION

        version.aprobar()
        version.save()
        assert version.estado == Version.Estado.APROBADA
        assert version.fecha_aprobacion is not None

    def test_flujo_rechazo_revisor(self, version):
        """Flujo: borrador → enviada → en_revision → rechazada."""
        version.enviar()
        version.save()
        version.tomar_revision()
        version.save()
        version.rechazar()
        version.save()
        assert version.estado == Version.Estado.RECHAZADA


# ============================================================
# TESTS DE ENTREGA TARDÍA
# ============================================================

@pytest.mark.django_db
class TestEntregaTardia:
    """Tests de la marca de entrega tardía."""

    def test_entrega_en_plazo(self, materia, profesor):
        """Entrega dentro del plazo no es marcada como tardía."""
        inst = InstanciaPresentacion.objects.create(
            nombre='Test En Plazo',
            anio_academico=2026,
            periodo=InstanciaPresentacion.Periodo.ANUAL,
            fecha_apertura=date.today() - timedelta(days=10),
            fecha_limite=date.today() + timedelta(days=10),
        )
        inst.carreras.add(materia.carrera)

        planif, _ = Planificacion.objects.get_or_create(materia=materia, profesor=profesor, instancia=inst)
        archivo = SimpleUploadedFile('test.docx', b'x')
        version = Version.objects.create(planificacion=planif, numero=1, archivo=archivo)

        version.enviar()
        version.save()

        assert version.entrega_tardia is False
        assert version.dias_atraso == 0

    def test_entrega_tardia(self, materia, profesor):
        """Entrega fuera del plazo es marcada como tardía."""
        inst = InstanciaPresentacion.objects.create(
            nombre='Test Tardía',
            anio_academico=2024,
            periodo=InstanciaPresentacion.Periodo.ANUAL,
            fecha_apertura=date(2024, 3, 1),
            fecha_limite=date(2024, 3, 31),
            estado=InstanciaPresentacion.Estado.ABIERTA,
        )
        inst.carreras.add(materia.carrera)

        planif, _ = Planificacion.objects.get_or_create(materia=materia, profesor=profesor, instancia=inst)
        archivo = SimpleUploadedFile('test.docx', b'x')
        version = Version.objects.create(planificacion=planif, numero=1, archivo=archivo)

        version.enviar()
        version.save()

        assert version.entrega_tardia is True
        assert version.dias_atraso > 0


# ============================================================
# TESTS DE VISTAS
# ============================================================

@pytest.mark.django_db
class TestPlanificacionViews:
    """Tests de las vistas de planificaciones."""

    def test_cargar_requiere_login(self, instancia, materia):
        """Vista cargar redirige si no está autenticado."""
        client = Client()
        response = client.get(f'/planificaciones/cargar/{instancia.pk}/{materia.pk}/')
        assert response.status_code == 302

    def test_detalle_requiere_login(self, planificacion):
        """Vista detalle redirige si no está autenticado."""
        client = Client()
        response = client.get(f'/planificaciones/{planificacion.pk}/')
        assert response.status_code == 302

    def test_profesor_accede_a_cargar(self, instancia, materia, usuario_profesor):
        """Profesor titular puede acceder a la vista de cargar."""
        client = Client()
        client.force_login(usuario_profesor)
        response = client.get(f'/planificaciones/cargar/{instancia.pk}/{materia.pk}/')
        assert response.status_code == 200

    def test_subir_archivo_crea_version(self, instancia, materia, usuario_profesor):
        """POST con archivo Word válido crea un borrador."""
        client = Client()
        client.force_login(usuario_profesor)

        archivo = SimpleUploadedFile(
            'plan.docx',
            doc_completo().read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response = client.post(
            f'/planificaciones/cargar/{instancia.pk}/{materia.pk}/',
            {'archivo': archivo}
        )
        assert response.status_code == 302
        # Debe haber creado una versión
        planif = Planificacion.objects.get(materia=materia, instancia=instancia)
        assert planif.versiones.count() == 1
        assert planif.ultima_version.estado == Version.Estado.BORRADOR

    def test_enviar_version_valida(self, planificacion, usuario_profesor):
        """Enviar una versión con doc válido cambia estado a ENVIADA."""
        client = Client()
        client.force_login(usuario_profesor)

        archivo = SimpleUploadedFile(
            'plan.docx',
            doc_completo().read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        version = Version.objects.create(
            planificacion=planificacion,
            numero=1,
            archivo=archivo
        )

        response = client.get(f'/planificaciones/enviar/{version.pk}/')
        version.refresh_from_db()
        assert response.status_code == 302
        assert version.estado == Version.Estado.ENVIADA

    def test_enviar_version_invalida_rechaza_auto(self, planificacion, usuario_profesor):
        """Enviar una versión con doc inválido cambia estado a RECHAZADA_AUTO."""
        client = Client()
        client.force_login(usuario_profesor)

        archivo = SimpleUploadedFile(
            'plan_incompleto.docx',
            doc_incompleto().read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        version = Version.objects.create(
            planificacion=planificacion,
            numero=1,
            archivo=archivo
        )

        response = client.get(f'/planificaciones/enviar/{version.pk}/')
        version.refresh_from_db()
        assert response.status_code == 302
        assert version.estado == Version.Estado.RECHAZADA_AUTO
        assert len(version.campos_faltantes) > 0
