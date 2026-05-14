"""
Comando para insertar datos de prueba coherentes.

Crea:
  - 2 instituciones (ICES, UCSE)
  - 2 carreras con coordinadores
  - 7 profesores con sus materias
  - Instancias 2024, 2025 (cerradas) y 2026 (anuales abiertas, 1er/2do cuat programadas)
  - Planificaciones con versiones en distintos estados

Uso:
    docker-compose exec web python manage.py seed_data
    docker-compose exec web python manage.py seed_data --flush   # borra todo primero
"""
import io
from datetime import date, timedelta

from django.core.management.base import BaseCommand
from django.db import transaction
from docx import Document


# ─────────────────────────── helpers ────────────────────────────

def make_docx(*extra_lines) -> bytes:
    """Genera un .docx mínimo con los 7 campos obligatorios."""
    doc = Document()
    doc.add_heading('Planificación de la Asignatura', level=1)
    campos = [
        'Propósito',
        'Fundamentación',
        'Contenidos Mínimos',
        'Metodología',
        'Requisitos para regularizar y promocionar',
        'Criterios y formatos de evaluación',
        'Bibliografía',
    ]
    for campo in campos:
        doc.add_heading(campo, level=2)
        doc.add_paragraph('Contenido de ejemplo para el campo.')
    for line in extra_lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_docx_incompleto() -> bytes:
    """Genera un .docx sin bibliografía (para simular rechazo automático)."""
    doc = Document()
    doc.add_heading('Planificación', level=1)
    for campo in ['Propósito', 'Fundamentación', 'Contenidos Mínimos',
                  'Metodología', 'Evaluación']:
        doc.add_heading(campo, level=2)
        doc.add_paragraph('Contenido de ejemplo.')
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


TODAY = date.today()


class Command(BaseCommand):
    help = 'Inserta datos de prueba coherentes para el sistema de planificaciones.'

    def add_arguments(self, parser):
        parser.add_argument(
            '--flush', action='store_true',
            help='Elimina todos los datos existentes antes de insertar.'
        )
        parser.add_argument(
            '--reset', action='store_true',
            help='Vacía la base y deja solo admin, moderadora y coordinadores (sin datos de prueba).'
        )
        parser.add_argument(
            '--catalogo', action='store_true',
            help='Vacía la base y crea moderadora, coordinadores, carreras, profesores y materias. Las instancias las creás vos.'
        )

    @transaction.atomic
    def handle(self, *args, **options):
        if options['reset']:
            self._flush()
            self.stdout.write('\n▶  Creando usuarios staff...')
            self._create_staff()
            self.stdout.write(self.style.SUCCESS('\n✅  Base reseteada. Solo quedan admin y moderadora. Creá coordinadores desde el Catálogo.\n'))
            return

        if options['catalogo']:
            self._flush()
            self.stdout.write('\n▶  Creando moderadora...')
            self._create_staff()
            self.stdout.write('▶  Creando catálogo (coordinadores + carreras + profesores + materias)...')
            self._create_catalogo()
            self.stdout.write(self.style.SUCCESS('\n✅  Catálogo listo. Ahora creá las instancias desde el sistema.\n'))
            self._print_catalogo_resumen()
            return

        if options['flush']:
            self._flush()

        self.stdout.write('\n▶  Creando usuarios...')
        users = self._create_users()

        self.stdout.write('▶  Creando catálogos...')
        catalogos = self._create_catalogos(users)

        self.stdout.write('▶  Creando instancias...')
        instancias = self._create_instancias(catalogos)

        self.stdout.write('▶  Creando planificaciones...')
        self._create_planificaciones(users, catalogos, instancias)

        self.stdout.write(self.style.SUCCESS('\n✅  Datos de prueba insertados correctamente.\n'))
        self._print_resumen(users)

    # ──────────────────────────────────────────────────────────────

    def _flush(self):
        from apps.revisiones.models import Revision, VistoBueno
        from apps.planificaciones.models import Version, Planificacion
        from apps.instancias.models import InstanciaPresentacion
        from apps.catalogos.models import Materia, Profesor, Carrera, Plantilla, MaterialApoyo, Institucion
        from apps.usuarios.models import Usuario

        self.stdout.write('⚠  Eliminando datos existentes...')
        Revision.objects.all().delete()
        VistoBueno.objects.all().delete()
        Version.objects.all().delete()
        Planificacion.objects.all().delete()
        InstanciaPresentacion.objects.all().delete()
        Materia.objects.all().delete()
        Profesor.objects.all().delete()
        Carrera.objects.all().delete()
        Plantilla.objects.all().delete()
        MaterialApoyo.objects.all().delete()
        Institucion.objects.all().delete()
        Usuario.objects.filter(is_superuser=False).delete()

    # ──────────────────────────────────────────────────────────────

    def _create_staff(self):
        """Crea solo la moderadora (sin coordinadores, para que los cree via ABM)."""
        from apps.usuarios.models import Usuario

        u, created = Usuario.objects.get_or_create(
            email='moderadora@ices.edu',
            defaults={'rol': 'moderadora', 'nombre_completo': 'Ana García'}
        )
        if created:
            u.set_password('mod123')
            u.save()

        self.stdout.write('\n Credenciales:')
        self.stdout.write('  moderadora@ices.edu  mod123')

    # ──────────────────────────────────────────────────────────────

    def _create_catalogo(self):
        """Crea coordinadores, carreras, profesores y materias sin instancias."""
        from apps.usuarios.models import Usuario
        from apps.catalogos.models import Institucion, Carrera, Profesor, Materia

        def usuario(email, password, rol, nombre):
            u, created = Usuario.objects.get_or_create(
                email=email, defaults={'rol': rol, 'nombre_completo': nombre}
            )
            if created:
                u.set_password(password)
                u.save()
            return u

        mod = usuario('moderadora@ices.edu', 'mod123', 'moderadora', 'Ana García')

        coord_sis  = usuario('coord.sistemas@ices.edu',     'coord123', 'coordinador', 'Carlos Rodríguez')
        coord_cont = usuario('coord.contabilidad@ices.edu', 'coord123', 'coordinador', 'Laura Martínez')

        u_perez  = usuario('perez@ices.edu',  'prof123', 'profesor', 'Roberto Pérez')
        u_gomez  = usuario('gomez@ices.edu',  'prof123', 'profesor', 'Claudia Gómez')
        u_silva  = usuario('silva@ices.edu',  'prof123', 'profesor', 'Marcelo Silva')
        u_torres = usuario('torres@ices.edu', 'prof123', 'profesor', 'Patricia Torres')

        ices, _ = Institucion.objects.get_or_create(
            codigo='ICES',
            defaults={'nombre': 'Instituto Cooperativo de Enseñanza Superior'}
        )

        sistemas, _ = Carrera.objects.get_or_create(
            nombre='Ingeniería en Sistemas', institucion=ices,
            defaults={'coordinador': coord_sis}
        )
        contabilidad, _ = Carrera.objects.get_or_create(
            nombre='Licenciatura en Administración', institucion=ices,
            defaults={'coordinador': coord_cont}
        )

        def prof(usuario_obj):
            p, _ = Profesor.objects.get_or_create(
                usuario=usuario_obj,
                defaults={'institucion': ices}
            )
            return p

        p_perez  = prof(u_perez)
        p_gomez  = prof(u_gomez)
        p_silva  = prof(u_silva)
        p_torres = prof(u_torres)

        def mat(nombre, carrera, anio, regimen, titular=None):
            Materia.objects.get_or_create(
                nombre=nombre, carrera=carrera,
                defaults={'anio_cursado': anio, 'regimen': regimen,
                          'profesor_titular': titular}
            )

        # Ingeniería en Sistemas
        mat('Programación I',          sistemas,     1, 'anual', p_perez)
        mat('Programación II',         sistemas,     2, 'anual', p_perez)
        mat('Base de Datos',           sistemas,     2, '1cuat', p_gomez)
        mat('Redes y Comunicaciones',  sistemas,     3, '1cuat', p_silva)
        mat('Ingeniería de Software',  sistemas,     3, '2cuat', p_gomez)
        mat('Sistemas Operativos',     sistemas,     3, 'anual', p_silva)
        mat('Arquitectura de Computadoras', sistemas, 2, '2cuat', p_silva)
        mat('Matemática Discreta',     sistemas,     1, '1cuat', p_perez)

        # Licenciatura en Administración
        mat('Contabilidad I',          contabilidad, 1, 'anual',  p_torres)
        mat('Contabilidad II',         contabilidad, 2, 'anual',  p_torres)
        mat('Matemática Financiera',   contabilidad, 2, '1cuat',  p_torres)
        mat('Administración General',  contabilidad, 1, 'anual',  p_perez)
        mat('Economía',                contabilidad, 1, '2cuat',  p_perez)
        mat('Derecho Laboral',         contabilidad, 3, '1cuat',  None)
        mat('Estadística Aplicada',    contabilidad, 2, '2cuat',  p_gomez)

    def _print_catalogo_resumen(self):
        self.stdout.write('\n Credenciales:')
        self.stdout.write('  moderadora@ices.edu         mod123')
        self.stdout.write('  coord.sistemas@ices.edu     coord123')
        self.stdout.write('  coord.contabilidad@ices.edu coord123')
        self.stdout.write('  perez@ices.edu              prof123')
        self.stdout.write('  gomez@ices.edu              prof123')
        self.stdout.write('  silva@ices.edu              prof123')
        self.stdout.write('  torres@ices.edu             prof123')

    # ──────────────────────────────────────────────────────────────

    def _create_users(self):
        from apps.usuarios.models import Usuario

        def get_or_create(email, password, rol, nombre):
            u, created = Usuario.objects.get_or_create(
                email=email,
                defaults={'rol': rol, 'nombre_completo': nombre}
            )
            if created:
                u.set_password(password)
                u.save()
            return u

        return {
            'moderadora': get_or_create(
                'moderadora@ices.edu', 'mod123', 'moderadora', 'Ana García'),
            'coord_sistemas': get_or_create(
                'coord.sistemas@ices.edu', 'coord123', 'coordinador', 'Carlos Rodríguez'),
            'coord_contabilidad': get_or_create(
                'coord.contabilidad@ices.edu', 'coord123', 'coordinador', 'Laura Martínez'),
            'prof_perez': get_or_create(
                'perez@ices.edu', 'prof123', 'profesor', 'Roberto Pérez'),
            'prof_gomez': get_or_create(
                'gomez@ices.edu', 'prof123', 'profesor', 'Claudia Gómez'),
            'prof_silva': get_or_create(
                'silva@ices.edu', 'prof123', 'profesor', 'Marcelo Silva'),
            'prof_torres': get_or_create(
                'torres@ices.edu', 'prof123', 'profesor', 'Patricia Torres'),
            'alumno': get_or_create(
                'alumno@ices.edu', 'alum123', 'alumno', 'Pedro Alumno'),
        }

    # ──────────────────────────────────────────────────────────────

    def _create_catalogos(self, users):
        from apps.catalogos.models import Institucion, Carrera, Profesor, Materia

        ices, _ = Institucion.objects.get_or_create(
            codigo='ICES',
            defaults={'nombre': 'Instituto Cooperativo de Enseñanza Superior'}
        )
        ucse, _ = Institucion.objects.get_or_create(
            codigo='UCSE',
            defaults={'nombre': 'Universidad Católica de Santiago del Estero'}
        )

        sistemas, _ = Carrera.objects.get_or_create(
            nombre='Ingeniería en Sistemas',
            institucion=ices,
            defaults={'coordinador': users['coord_sistemas']}
        )
        contabilidad, _ = Carrera.objects.get_or_create(
            nombre='Licenciatura en Administración',
            institucion=ices,
            defaults={'coordinador': users['coord_contabilidad']}
        )

        def get_or_create_prof(usuario, institucion):
            p, _ = Profesor.objects.get_or_create(
                usuario=usuario,
                defaults={'institucion': institucion}
            )
            return p

        prof_perez   = get_or_create_prof(users['prof_perez'],  ices)
        prof_gomez   = get_or_create_prof(users['prof_gomez'],  ices)
        prof_silva   = get_or_create_prof(users['prof_silva'],  ices)
        prof_torres  = get_or_create_prof(users['prof_torres'], ices)

        def mat(nombre, carrera, anio, regimen, prof=None):
            m, _ = Materia.objects.get_or_create(
                nombre=nombre, carrera=carrera,
                defaults={'anio_cursado': anio, 'regimen': regimen,
                          'profesor_titular': prof}
            )
            return m

        # Sistemas
        prog1   = mat('Programación I',        sistemas, 1, 'anual',  prof_perez)
        prog2   = mat('Programación II',       sistemas, 2, 'anual',  prof_perez)
        bdatos  = mat('Base de Datos',         sistemas, 2, '1cuat',  prof_gomez)
        redes   = mat('Redes y Comunicaciones',sistemas, 3, '1cuat',  prof_silva)
        iso     = mat('Ingeniería de Software',sistemas, 3, '2cuat',  prof_gomez)
        so      = mat('Sistemas Operativos',   sistemas, 3, 'anual',  prof_silva)

        # Administración
        cont1   = mat('Contabilidad I',        contabilidad, 1, 'anual',  prof_torres)
        mat_fin = mat('Matemática Financiera', contabilidad, 2, '1cuat',  prof_torres)
        admin_g = mat('Administración General',contabilidad, 1, 'anual',  prof_perez)

        return {
            'ices': ices, 'ucse': ucse,
            'sistemas': sistemas, 'contabilidad': contabilidad,
            'prof_perez': prof_perez, 'prof_gomez': prof_gomez,
            'prof_silva': prof_silva, 'prof_torres': prof_torres,
            'prog1': prog1, 'prog2': prog2, 'bdatos': bdatos,
            'redes': redes, 'iso': iso, 'so': so,
            'cont1': cont1, 'mat_fin': mat_fin, 'admin_g': admin_g,
        }

    # ──────────────────────────────────────────────────────────────

    def _create_instancias(self, cat):
        from apps.instancias.models import InstanciaPresentacion as Inst

        sistemas     = cat['sistemas']
        contabilidad = cat['contabilidad']

        def inst(nombre, anio, periodo, apertura, limite, cerrada=False, carreras=None):
            obj, _ = Inst.objects.get_or_create(
                nombre=nombre, anio_academico=anio,
                defaults={
                    'periodo': periodo,
                    'solo_regimen': periodo,  # por defecto, filtrar por el mismo régimen
                    'fecha_apertura': apertura,
                    'fecha_limite': limite,
                    'estado': 'cerrada' if cerrada else 'programada',
                }
            )
            if carreras:
                for c in carreras:
                    obj.carreras.add(c)
            return obj

        ambas = [sistemas, contabilidad]

        # ── 2026 ──
        anual_26 = inst(
            'Anuales', 2026, 'anual',
            apertura=TODAY - timedelta(days=7),
            limite=TODAY + timedelta(days=21),
            carreras=ambas,
        )
        primer_26 = inst(
            '1° Cuatrimestre', 2026, '1cuat',
            apertura=TODAY - timedelta(days=3),
            limite=TODAY + timedelta(days=14),
            carreras=ambas,
        )
        segundo_26 = inst(
            '2° Cuatrimestre', 2026, '2cuat',
            apertura=TODAY + timedelta(days=60),
            limite=TODAY + timedelta(days=90),
            carreras=ambas,
        )

        # ── 2025 ──
        anual_25 = inst(
            'Anuales', 2025, 'anual',
            apertura=date(2025, 3, 1), limite=date(2025, 3, 31),
            cerrada=True, carreras=ambas,
        )
        primer_25 = inst(
            '1° Cuatrimestre', 2025, '1cuat',
            apertura=date(2025, 2, 10), limite=date(2025, 2, 28),
            cerrada=True, carreras=ambas,
        )
        segundo_25 = inst(
            '2° Cuatrimestre', 2025, '2cuat',
            apertura=date(2025, 6, 1), limite=date(2025, 6, 20),
            cerrada=True, carreras=ambas,
        )

        # ── 2024 ──
        anual_24 = inst(
            'Anuales', 2024, 'anual',
            apertura=date(2024, 3, 1), limite=date(2024, 3, 31),
            cerrada=True, carreras=ambas,
        )
        primer_24 = inst(
            '1° Cuatrimestre', 2024, '1cuat',
            apertura=date(2024, 2, 10), limite=date(2024, 2, 28),
            cerrada=True, carreras=ambas,
        )

        return {
            'anual_26': anual_26, 'primer_26': primer_26, 'segundo_26': segundo_26,
            'anual_25': anual_25, 'primer_25': primer_25, 'segundo_25': segundo_25,
            'anual_24': anual_24, 'primer_24': primer_24,
        }

    # ──────────────────────────────────────────────────────────────

    def _create_planificaciones(self, users, cat, inst):
        from django.core.files.base import ContentFile
        from apps.planificaciones.models import Planificacion, Version
        from apps.revisiones.models import Revision, VistoBueno

        def make_version(planif, numero, estado, tardio=False, incompleto=False,
                         dias_atraso=0, campos_faltantes=None):
            if Version.objects.filter(planificacion=planif, numero=numero).exists():
                return Version.objects.get(planificacion=planif, numero=numero)
            contenido = make_docx_incompleto() if incompleto else make_docx(f'Versión {numero}')
            v = Version(planificacion=planif, numero=numero,
                        entrega_tardia=tardio, dias_atraso=dias_atraso)
            v.archivo.save(f'plan_{planif.pk}_v{numero}.docx', ContentFile(contenido), save=False)
            if campos_faltantes:
                v.campos_faltantes = campos_faltantes
            v.save()
            # Aplicar estado FSM
            if estado in ('en_revision', 'rechazada', 'rechazada_auto',
                          'aprobada', 'oficial', 'reemplazada'):
                v.estado = estado
                from django.utils import timezone
                v.fecha_envio = timezone.now()
                if estado in ('aprobada', 'oficial'):
                    v.fecha_aprobacion = timezone.now()
                v.save()
            return v

        def planif(materia, profesor, instancia):
            p, _ = Planificacion.objects.get_or_create(
                materia=materia, profesor=profesor, instancia=instancia
            )
            return p

        def visto(version, usuario, rol):
            VistoBueno.objects.get_or_create(
                version=version, rol=rol,
                defaults={'usuario': usuario}
            )

        # ── Instancia anual 2026 (abierta) ──────────────────────────

        # Prog I – en revisión (recién enviada)
        p = planif(cat['prog1'], cat['prof_perez'], inst['anual_26'])
        make_version(p, 1, 'en_revision')

        # Prog II – en revisión, solo visto de moderadora
        p = planif(cat['prog2'], cat['prof_perez'], inst['anual_26'])
        v = make_version(p, 1, 'en_revision')
        visto(v, users['moderadora'], 'moderadora')
        Revision.objects.get_or_create(
            planificacion=p, version=v, tipo='aprobar',
            defaults={'usuario': users['moderadora'], 'observaciones': 'Visto bueno de moderadora'})

        # Base de Datos – rechazada auto (doc incompleto) + nueva v2 en borrador
        p = planif(cat['bdatos'], cat['prof_gomez'], inst['anual_26'])
        make_version(p, 1, 'rechazada_auto', incompleto=True,
                     campos_faltantes=['bibliografia', 'requisitos'])
        make_version(p, 2, 'borrador')

        # Redes – rechazada por revisor con observaciones + v2 en revisión
        p = planif(cat['redes'], cat['prof_silva'], inst['anual_26'])
        v1 = make_version(p, 1, 'rechazada')
        Revision.objects.get_or_create(
            planificacion=p, version=v1, tipo='rechazar',
            defaults={'usuario': users['moderadora'],
                      'observaciones': 'La bibliografía no está actualizada. Incluir fuentes del 2020 en adelante.'})
        make_version(p, 2, 'en_revision')

        # Ing. Software – borrador (todavía no envió)
        p = planif(cat['iso'], cat['prof_gomez'], inst['anual_26'])
        make_version(p, 1, 'borrador')

        # Sistemas Operativos – tardía, en revisión
        p = planif(cat['so'], cat['prof_silva'], inst['anual_26'])
        make_version(p, 1, 'en_revision', tardio=True, dias_atraso=4)

        # Contabilidad I – oficial (aprobada con doble visto)
        p = planif(cat['cont1'], cat['prof_torres'], inst['anual_26'])
        v = make_version(p, 1, 'oficial')
        p.version_oficial = v
        p.save()
        visto(v, users['moderadora'], 'moderadora')
        visto(v, users['coord_contabilidad'], 'coordinador')

        # ── Instancia 1er cuatrimestre 2026 (abierta) ────────────────

        p = planif(cat['bdatos'], cat['prof_gomez'], inst['primer_26'])
        make_version(p, 1, 'en_revision')

        p = planif(cat['redes'], cat['prof_silva'], inst['primer_26'])
        make_version(p, 1, 'borrador')

        p = planif(cat['mat_fin'], cat['prof_torres'], inst['primer_26'])
        make_version(p, 1, 'oficial')
        v = Version.objects.filter(planificacion=p).first()
        p.version_oficial = v
        p.save()
        visto(v, users['moderadora'], 'moderadora')
        visto(v, users['coord_contabilidad'], 'coordinador')

        # ── Instancias 2025 (cerradas, con planificaciones oficiales) ──

        for mat_obj, prof_obj, coord in [
            (cat['prog1'],  cat['prof_perez'],  users['coord_sistemas']),
            (cat['prog2'],  cat['prof_perez'],  users['coord_sistemas']),
            (cat['bdatos'], cat['prof_gomez'],  users['coord_sistemas']),
            (cat['cont1'],  cat['prof_torres'], users['coord_contabilidad']),
            (cat['mat_fin'],cat['prof_torres'], users['coord_contabilidad']),
        ]:
            p = planif(mat_obj, prof_obj, inst['anual_25'])
            if not Version.objects.filter(planificacion=p).exists():
                v = make_version(p, 1, 'oficial')
                p.version_oficial = v
                p.save()
                visto(v, users['moderadora'], 'moderadora')
                visto(v, coord, 'coordinador')

        # ── Instancias 2024 (cerradas, con planificaciones oficiales) ──

        for mat_obj, prof_obj, coord in [
            (cat['prog1'],  cat['prof_perez'],  users['coord_sistemas']),
            (cat['bdatos'], cat['prof_gomez'],  users['coord_sistemas']),
            (cat['cont1'],  cat['prof_torres'], users['coord_contabilidad']),
        ]:
            p = planif(mat_obj, prof_obj, inst['anual_24'])
            if not Version.objects.filter(planificacion=p).exists():
                v = make_version(p, 1, 'oficial')
                p.version_oficial = v
                p.save()
                visto(v, users['moderadora'], 'moderadora')
                visto(v, coord, 'coordinador')

    # ──────────────────────────────────────────────────────────────

    def _print_resumen(self, users):
        self.stdout.write('\n──────────────────────────────────────────')
        self.stdout.write(' Usuarios de prueba:')
        self.stdout.write('──────────────────────────────────────────')
        creds = [
            ('Moderadora',             'moderadora@ices.edu',       'mod123'),
            ('Coordinador Sistemas',   'coord.sistemas@ices.edu',   'coord123'),
            ('Coordinador Administ.', 'coord.contabilidad@ices.edu','coord123'),
            ('Profesor Pérez',         'perez@ices.edu',            'prof123'),
            ('Profesor Gómez',         'gomez@ices.edu',            'prof123'),
            ('Profesor Silva',         'silva@ices.edu',            'prof123'),
            ('Profesor Torres',        'torres@ices.edu',           'prof123'),
            ('Alumno',                 'alumno@ices.edu',           'alum123'),
        ]
        for rol, email, pwd in creds:
            self.stdout.write(f'  {rol:<28} {email:<35} {pwd}')
        self.stdout.write('──────────────────────────────────────────\n')
